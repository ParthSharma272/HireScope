"""
DOCX Template Generator
AI-powered resume template generation with industry-specific optimization
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Optional


class TemplateGenerator:
    """Generate professionally formatted DOCX templates with AI optimization"""
    
    def __init__(self):
        self.style_config = {
            'tech': {
                'accent_color': RGBColor(147, 51, 234),  # Purple
                'font': 'Calibri',
                'description': 'Modern technical design'
            },
            'finance': {
                'accent_color': RGBColor(37, 99, 235),  # Blue
                'font': 'Georgia',
                'description': 'Conservative professional'
            },
            'healthcare': {
                'accent_color': RGBColor(22, 163, 74),  # Green
                'font': 'Arial',
                'description': 'Clean and trustworthy'
            },
            'marketing': {
                'accent_color': RGBColor(236, 72, 153),  # Pink
                'font': 'Helvetica',
                'description': 'Creative and bold'
            }
        }
        
        # Industry-specific section templates
        self.section_templates = {
            'tech': {
                'order': ['Professional Summary', 'Technical Skills', 'Experience', 'Projects', 'Education'],
                'required': ['Technical Skills', 'Experience', 'Education'],
                'sample_content': {
                    'Professional Summary': 'Results-driven Software Engineer with [X] years of experience in full-stack development, cloud architecture, and scalable system design. Proven track record of delivering high-impact solutions.',
                    'Technical Skills': {
                        'Languages': ['Python', 'JavaScript', 'TypeScript', 'SQL', 'Java'],
                        'Frameworks': ['React', 'Django', 'FastAPI', 'Node.js', 'Express'],
                        'Tools & Platforms': ['Docker', 'Kubernetes', 'AWS', 'Git', 'Jenkins']
                    },
                    'Experience': [
                        {
                            'company': 'Tech Company Inc.',
                            'title': 'Senior Software Engineer',
                            'location': 'San Francisco, CA',
                            'dates': 'Jan 2020 - Present',
                            'bullets': [
                                'Architected microservices platform processing 2M+ requests/day, reducing API latency by 40%',
                                'Led team of 5 engineers in developing real-time analytics dashboard using React and FastAPI',
                                'Implemented CI/CD pipeline reducing deployment time from 2 hours to 15 minutes'
                            ]
                        }
                    ]
                }
            },
            'finance': {
                'order': ['Professional Summary', 'Core Competencies', 'Experience', 'Certifications', 'Education'],
                'required': ['Professional Summary', 'Experience', 'Education'],
                'sample_content': {
                    'Professional Summary': 'Strategic Financial Analyst with [X] years of experience in portfolio management, risk assessment, and financial modeling. CPA-certified with proven expertise in driving ROI.',
                    'Core Competencies': ['Financial Modeling', 'Portfolio Management', 'Risk Analysis', 'Excel/VBA', 'Bloomberg Terminal', 'Regulatory Compliance'],
                    'Experience': [
                        {
                            'company': 'Investment Firm LLC',
                            'title': 'Senior Financial Analyst',
                            'location': 'New York, NY',
                            'dates': 'Mar 2018 - Present',
                            'bullets': [
                                'Managed $50M portfolio with 18% YoY returns, exceeding benchmark by 6%',
                                'Developed financial models reducing forecasting error by 25% through advanced Excel/VBA automation',
                                'Conducted risk assessments for 100+ investment opportunities, maintaining 95% accuracy rate'
                            ]
                        }
                    ]
                }
            },
            'healthcare': {
                'order': ['Licensure & Certifications', 'Clinical Experience', 'Skills', 'Education'],
                'required': ['Licensure & Certifications', 'Clinical Experience', 'Education'],
                'sample_content': {
                    'Licensure & Certifications': ['RN License (State)', 'BLS Certification', 'ACLS Certification', 'HIPAA Compliance Certified'],
                    'Clinical Experience': [
                        {
                            'company': 'City General Hospital',
                            'title': 'Registered Nurse',
                            'location': 'Chicago, IL',
                            'dates': 'Jun 2019 - Present',
                            'bullets': [
                                'Provided patient care for 15-20 patients per shift in high-acuity medical-surgical unit',
                                'Improved patient satisfaction scores by 25% through compassionate care and clear communication',
                                'Collaborated with interdisciplinary team to implement new EMR system, improving documentation efficiency by 30%'
                            ]
                        }
                    ]
                }
            },
            'marketing': {
                'order': ['Professional Summary', 'Key Achievements', 'Experience', 'Skills & Tools', 'Education'],
                'required': ['Professional Summary', 'Experience', 'Education'],
                'sample_content': {
                    'Professional Summary': 'Data-driven Marketing Manager with [X] years of experience in digital campaigns, brand strategy, and ROI optimization. Proven success in scaling revenue through innovative marketing.',
                    'Key Achievements': [
                        'Led campaign generating $2.5M revenue with 320% ROI and 45% conversion rate',
                        'Grew organic traffic by 150% in 12 months through SEO optimization',
                        'Managed $500K digital ad budget across Google Ads, Facebook, and LinkedIn'
                    ],
                    'Experience': [
                        {
                            'company': 'Digital Marketing Agency',
                            'title': 'Marketing Manager',
                            'location': 'Austin, TX',
                            'dates': 'Feb 2019 - Present',
                            'bullets': [
                                'Spearheaded multi-channel campaigns reaching 2M+ users, increasing brand awareness by 85%',
                                'Optimized email marketing strategy achieving 28% open rate and 8% CTR, 2x industry average',
                                'Managed team of 4 content creators producing 50+ pieces of engaging content monthly'
                            ]
                        }
                    ]
                }
            }
        }
    
    def generate_template(
        self,
        industry: str,
        sections: Optional[List[str]] = None,
        user_content: Optional[Dict] = None,
        ai_enhancements: Optional[Dict] = None
    ) -> bytes:
        """
        Generate ATS-optimized DOCX template
        
        Args:
            industry: Industry type (tech, finance, healthcare, marketing)
            sections: List of sections to include (uses defaults if None)
            user_content: Optional user data to populate template
            ai_enhancements: Optional AI-enhanced content
            
        Returns:
            bytes: DOCX file as bytes
        """
        # Validate industry
        if industry not in self.style_config:
            industry = 'tech'  # Default fallback
        
        # Get style and content
        style = self.style_config[industry]
        template_data = self.section_templates[industry]
        
        # Use provided sections or default order
        sections_to_include = sections or template_data['order']
        
        # Create document
        doc = Document()
        
        # Set margins (ATS-friendly: 0.75 inch)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add header with name
        name = (user_content or {}).get('name', 'YOUR NAME')
        self._add_name_header(doc, name, style)
        
        # Add contact info
        contact = (user_content or {}).get('contact', {})
        self._add_contact_info(doc, contact)
        
        # Add sections dynamically
        for section_name in sections_to_include:
            content = None
            if user_content and section_name.lower().replace(' ', '_') in user_content:
                content = user_content[section_name.lower().replace(' ', '_')]
            else:
                content = template_data['sample_content'].get(section_name, {})
            
            self._add_section(doc, section_name, style, content, ai_enhancements)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _add_name_header(self, doc: Document, name: str, style: Dict):
        """Add name header with accent color"""
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(name.upper())
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = style['accent_color']
        run.font.name = style['font']
    
    def _add_contact_info(self, doc: Document, contact: Dict):
        """Add contact information line"""
        email = contact.get('email', 'email@example.com')
        phone = contact.get('phone', '(123) 456-7890')
        location = contact.get('location', 'City, State')
        linkedin = contact.get('linkedin', '')
        github = contact.get('github', '')
        
        contact_parts = [email, phone, location]
        if linkedin:
            contact_parts.append(f'LinkedIn: {linkedin}')
        if github:
            contact_parts.append(f'GitHub: {github}')
        
        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_section(
        self,
        doc: Document,
        section_name: str,
        style: Dict,
        content: any,
        ai_enhancements: Optional[Dict]
    ):
        """Add a section with content"""
        # Section header
        header = doc.add_paragraph()
        run = header.add_run(section_name.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = style['accent_color']
        run.font.name = style['font']
        
        # Add horizontal line
        line_para = doc.add_paragraph()
        line_run = line_para.add_run('_' * 80)
        line_run.font.size = Pt(6)
        line_run.font.color.rgb = RGBColor(200, 200, 200)
        
        # Add content based on type
        if isinstance(content, str):
            # Simple text section
            p = doc.add_paragraph(content)
            p.runs[0].font.size = Pt(11)
        elif isinstance(content, list):
            if section_name.lower() in ['experience', 'clinical experience']:
                # Experience section with jobs
                for job in content:
                    self._add_experience_entry(doc, job, style)
            else:
                # Bullet list (certifications, achievements, etc.)
                for item in content:
                    p = doc.add_paragraph(item, style='List Bullet')
                    p.runs[0].font.size = Pt(10)
        elif isinstance(content, dict):
            if section_name.lower() in ['technical skills', 'core competencies', 'skills & tools', 'skills']:
                # Skills section with categories
                for category, skills in content.items():
                    p = doc.add_paragraph()
                    cat_run = p.add_run(f'{category}: ')
                    cat_run.font.bold = True
                    cat_run.font.size = Pt(10)
                    skills_run = p.add_run(', '.join(skills) if isinstance(skills, list) else skills)
                    skills_run.font.size = Pt(10)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_experience_entry(self, doc: Document, job: Dict, style: Dict):
        """Add an experience entry with bullets"""
        # Job header (company | title)
        job_para = doc.add_paragraph()
        company_title = job_para.add_run(f"{job.get('company', 'Company Name')} | {job.get('title', 'Job Title')}")
        company_title.font.bold = True
        company_title.font.size = Pt(11)
        
        # Location and dates
        loc_dates = doc.add_paragraph(f"{job.get('location', 'City, State')} | {job.get('dates', 'Month Year - Present')}")
        loc_dates.runs[0].font.size = Pt(10)
        loc_dates.runs[0].font.italic = True
        
        # Bullets
        bullets = job.get('bullets', [])
        for bullet in bullets:
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.runs[0].font.size = Pt(10)
        
        # Spacing between jobs
        doc.add_paragraph()
    
    def get_available_templates(self) -> List[Dict]:
        """Get list of available template types"""
        # Define RGB values directly for cleaner access
        color_map = {
            'tech': (147, 51, 234),     # Purple
            'finance': (37, 99, 235),    # Blue
            'healthcare': (22, 163, 74), # Green
            'marketing': (236, 72, 153)  # Pink
        }
        
        templates = []
        for industry, config in self.style_config.items():
            rgb = color_map[industry]
            
            templates.append({
                'id': industry,
                'name': industry.title(),
                'description': config['description'],
                'sections': self.section_templates[industry]['order'],
                'accent_color': f"rgb({rgb[0]}, {rgb[1]}, {rgb[2]})",
                'font': config['font']
            })
        return templates
